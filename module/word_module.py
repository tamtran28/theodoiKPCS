from docx import Document
import pandas as pd
import re


def clean_text(t):
    if not t:
        return ""
    return " ".join(t.replace("\n", " ").split()).strip()


def extract_right_block(text):
    """Tách text block bên phải thành dictionary"""

    out = {
        "trach_nhiem_thuc_hien": "",
        "trach_nhiem_quan_ly": "",
        "nguoi_phe_duyet": "",
        "muc_do_uu_tien": "",
        "ngay_hoan_thanh": "",
    }

    patterns = {
        "trach_nhiem_thuc_hien": r"Người chịu trách nhiệm thực hiện[:\- ]*(.+?)(?=Người phê duyệt|Ngày hoàn thành|$)",
        "nguoi_phe_duyet": r"Người phê duyệt[:\- ]*(.+?)(?=Ngày hoàn thành|$)",
        "muc_do_uu_tien": r"Mức độ ưu tiên hành động[:\- ]*(.+?)(?=Người chịu|Kế hoạch|$)",
        "ngay_hoan_thanh": r"Ngày hoàn thành[:\- ]*(.+)",
    }

    for key, pat in patterns.items():
        m = re.search(pat, text, flags=re.IGNORECASE | re.DOTALL)
        if m:
            out[key] = clean_text(m.group(1))

    return out


def split_phathien_ngnhan(text):
    """
    Tách cột "Phát hiện & Nguyên nhân" thành:
    - mô tả phát hiện
    - nguyên nhân
    """
    if not text:
        return "", ""

    # thường có dạng:
    # - Mô tả...
    # Nguyên nhân:
    parts = re.split(r"nguyên nhân[:\-]", text, flags=re.IGNORECASE)

    if len(parts) == 2:
        phathien = clean_text(parts[0])
        nguyennhan = clean_text(parts[1])
    else:
        phathien = clean_text(text)
        nguyennhan = ""

    return phathien, nguyennhan


def word_to_kiennghi(file):
    doc = Document(file)
    tables = []
    results = []

    # ======== 1) TÁCH BẢNG 5 CỘT ========
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [clean_text(c.text) for c in row.cells]
            rows.append(cells)

        # Bỏ header thừa
        rows = [r for r in rows if any(x.strip() for x in r)]

        # tìm header chính
        header = None
        for r in rows:
            if ("phát hiện" in " ".join(r).lower()
                    and "kiến nghị" in " ".join(r).lower()):
                header = r
                break

        if not header:
            continue

        idx = rows.index(header)
        data_rows = rows[idx + 1:]

        df = pd.DataFrame(data_rows, columns=header)
        tables.append(df)

    # ======== 2) TÁCH RIGHT-BLOCK TEXT ========
    full_text = "\n".join([p.text for p in doc.paragraphs])
    block_info = extract_right_block(full_text)

    # ======== 3) TÁCH PHÁT HIỆN – NGUYÊN NHÂN ========
    for df in tables:
        for _, row in df.iterrows():
            ph_text = row.get("Phát hiện và Nguyên nhân", "")

            mota, nguyennhan = split_phathien_ngnhan(ph_text)

            results.append({
                "Mô tả phát hiện": mota,
                "Nguyên nhân": nguyennhan,
                "Ảnh hưởng": row.get("Ảnh hưởng", ""),
                "Kiến nghị": row.get("Kiến nghị", ""),
                "Trách nhiệm thực hiện": block_info["trach_nhiem_thuc_hien"],
                "Trách nhiệm quản lý": block_info["trach_nhiem_quan_ly"],
                "Người phê duyệt": block_info["nguoi_phe_duyet"],
                "Mức độ ưu tiên": block_info["muc_do_uu_tien"],
                "Ngày hoàn thành": block_info["ngay_hoan_thanh"],
            })

    return pd.DataFrame(results)
