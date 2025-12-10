import pytesseract
import cv2
import numpy as np
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
from PyPDF2 import PdfReader
from io import BytesIO
import re
import tempfile


def read_word(file_obj):
    doc = Document(file_obj)
    return "\n".join(p.text for p in doc.paragraphs)


def read_pdf(file_obj):
    reader = PdfReader(file_obj)
    text = ""
    for page in reader.pages:
        try:
            txt = page.extract_text()
            if txt:
                text += txt + "\n"
        except:
            pass
    return text


def ocr_image(file_obj):
    img = Image.open(file_obj).convert("RGB")
    img_np = np.array(img)

    gray = cv2.cvtColor(img_np, cv2.COLOR_BGR2GRAY)
    gray = cv2.threshold(gray, 180, 255, cv2.THRESH_BINARY)[1]

    return pytesseract.image_to_string(gray, lang="vie")


def ocr_pdf(file_bytes: bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes)
        tmp.flush()
        pages = convert_from_path(tmp.name, dpi=250)

    result = ""
    for page in pages:
        img = np.array(page)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        gray = cv2.threshold(gray, 180, 255, cv2.THRESH_BINARY)[1]
        result += pytesseract.image_to_string(gray, lang="vie") + "\n"

    return result


def extract_kien_nghi(text: str):
    lower = text.lower()
    start = lower.find("kiến nghị")
    if start == -1:
        return []

    section = text[start:]

    parts = re.split(r"\n\s*\d+[\.\)]\s+", section)

    out = []
    for p in parts:
        p = p.strip()
        if len(p) <= 10:
            continue
        if p.lower().startswith("kiến nghị"):
            continue
        out.append(p)

    return out

# import pytesseract
# import cv2
# import numpy as np
# from PIL import Image
# from pdf2image import convert_from_path
# from docx import Document
# from PyPDF2 import PdfReader
# from io import BytesIO
# import re
# import tempfile


# def read_word(file_obj):
#     doc = Document(file_obj)
#     return "\n".join(p.text for p in doc.paragraphs)


# def read_pdf(file_obj):
#     reader = PdfReader(file_obj)
#     text = ""
#     for page in reader.pages:
#         try:
#             txt = page.extract_text()
#             if txt:
#                 text += txt + "\n"
#         except:
#             pass
#     return text


# def ocr_image(file_obj):
#     img = Image.open(file_obj).convert("RGB")
#     img_np = np.array(img)

#     gray = cv2.cvtColor(img_np, cv2.COLOR_BGR2GRAY)
#     gray = cv2.threshold(gray, 180, 255, cv2.THRESH_BINARY)[1]

#     return pytesseract.image_to_string(gray, lang="vie")


# def ocr_pdf(file_bytes: bytes):
#     with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
#         tmp.write(file_bytes)
#         tmp.flush()
#         pages = convert_from_path(tmp.name, dpi=250)

#     result = ""
#     for page in pages:
#         img = np.array(page)
#         gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
#         gray = cv2.threshold(gray, 180, 255, cv2.THRESH_BINARY)[1]
#         result += pytesseract.image_to_string(gray, lang="vie") + "\n"

#     return result


# def extract_kien_nghi(text: str):
#     lower = text.lower()
#     start = lower.find("kiến nghị")
#     if start == -1:
#         return []

#     section = text[start:]

#     parts = re.split(r"\n\s*\d+[\.\)]\s+", section)

#     out = []
#     for p in parts:
#         p = p.strip()
#         if len(p) <= 10:
#             continue
#         if p.lower().startswith("kiến nghị"):
#             continue
#         out.append(p)

#     return out
