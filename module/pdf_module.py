from docx import Document
import pandas as pd


def word_to_kiennghi(file):
    """
    Đọc bảng Word CHI TIẾT PHÁT HIỆN KIỂM TOÁN đúng format:
    - Lấy đúng header 5 cột
    - Bỏ dòng 'Chính sách, quy định...'
    - Bỏ dòng '1. ...' tiêu đề nhóm
    - Ghép các đoạn văn trong cell thành 1 dòng
    - Xử lý merge cell
    """

    doc = Document(file)
    tables = []

    for tbl in doc.tables:
        rows_out = []
        max_cols = 0

        # đọc toàn bộ cell
        for row in tbl.rows:
            cells = []
            for cell in row.cells:
                text = " ".join([p.text.strip() for p in cell.paragraphs]).strip()
                cells.append(text)

            # bỏ dòng hoàn toàn rỗng
            if all(c == "" for c in cells):
                continue

            rows_out.append(cells)
            max_cols = max(max_cols, len(cells))

        # chuẩn hóa số cột (bù thiếu do merge)
        clean_rows = []
        for r in rows_out:
            while len(r) < max_cols:
                r.append("")
            clean_rows.append(r)

        # tìm header thật (dòng có đủ 5 cột và chứa từ khóa)
        header = None
        for r in clean_rows:
            join = " ".join(r).lower()
            if "phát hiện" in join and "kiến nghị" in join:
                header = r
                break

        if header is None:
            continue

        # lấy index header
        header_idx = clean_rows.index(header)

        # bỏ phần trước header
        data_rows = clean_rows[header_idx + 1:]

        # loại bỏ dòng tiêu đề nhóm
        filtered = []
        for r in data_rows:
            row_text = " ".join(r).lower()

            if "chính sách" in row_text:
                continue
            if r[0].startswith("1.") or r[0].startswith("2.") or r[0].startswith("3."):
                # đây là dòng nhóm
                continue

            filtered.append(r)

        df = pd.DataFrame(filtered, columns=header)
        tables.append(df)

    return tables
